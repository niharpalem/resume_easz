import streamlit as st
import groq
import docx2txt
from prompts import quick_analyze_resume, in_depth_analyze_resume, enhance_resume, generate_formatted_resume_html
from flowchart import display_architecture_diagram
import docx
from io import BytesIO
from diff_match_patch import diff_match_patch
from fpdf import FPDF
from pdf2image import convert_from_path
import io

def update_docx(original_file, new_content):
    """
    Update the content of a docx file and return a BytesIO object with the updated file.
    """
    #to read the docx file
    doc = docx.Document(original_file)
    #clears the content of the docx file
    doc._body.clear_content()
    #adds the new content to the docx file
    doc.add_paragraph(new_content)
    #creates a BytesIO object to store the updated docx file (for temporary storage)
    doc_io = BytesIO()
    doc.save(doc_io)
    #moves the pointer back to start of the BytesIO object
    doc_io.seek(0)
    return doc_io

def show_diff(text1, text2):
    """
    Show the differences between two texts side by side
    """
    dmp = diff_match_patch()
    diffs = dmp.diff_main(text1, text2)
    dmp.diff_cleanupSemantic(diffs)
    
    html = []
    for (op, data) in diffs:
        text = data.replace('\n', '<br>')
        #0 for no change, 1 for insertion, -1 for deletion
        if op == 1:
            #highlight the inserted text with green background
            html.append(f'<span style="background-color: #d1fae5">{text}</span>')
        elif op == -1:
            #strike through the deleted text
            html.append(f'<span style="background-color: #fee2e2; text-decoration: line-through">{text}</span>')
        else:
            #no change
            html.append(text)
            
    return ''.join(html)


def main():
    st.set_page_config(
        layout="wide",
        page_title="Resume Enhancer",
        initial_sidebar_state="collapsed"
    )

    # Custom CSS with reduced text sizes
    st.markdown("""
        <style>
        .block-container {
            padding-top: 1.5rem;
            padding-bottom: 1.5rem;
            max-width: 1200px;
        }
        .stButton>button {
            background-color: #2563eb;
            color: white;
            border-radius: 0.375rem;
            padding: 0.75rem 1.5rem;
            border: none;
            box-shadow: 0 1px 2px rgba(0, 0, 0, 0.05);
            margin: 0.5rem;
            min-width: 200px;
            font-size: 0.875rem;
        }
        [data-testid="stFileUploader"] {
            border: 2px dashed #e5e7eb;
            border-radius: 0.5rem;
            padding: 0.875rem;
            min-height: 220px;
            font-size: 0.875rem;
        }
        .stTextArea>div>div {
            border-radius: 0.5rem;
            min-height: 220px !important;
            font-size: 0.875rem;
        }
        .stTextInput>div>div>input {
            border-radius: 0.5rem;
            font-size: 0.875rem;
        }
        .resume-html {
            padding: 1.5rem;
            max-width: 800px;
            margin: 0 auto;
            background: white;
            box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1);
            border-radius: 0.5rem;
            font-size: 0.875rem;
        }
        h1 {font-size: 3rem !important;  /* Adjust this value to increase the font size */
        } h2 {font-size: 1.5rem !important;  /* Adjust this value to increase the font size */
        h3, h4, h5, h6 {
            font-size: 80% !important;
        }
        p, li {
            font-size: 0.875rem !important;
        }
        </style>
    """, unsafe_allow_html=True)

    # Header with smaller text
    st.markdown("""
        <h1 style='text-align: center; font-size: 2.5rem; font-weight: 800; margin-bottom: 0.875rem;'>
            Resume Easz 
        </h1>
    """, unsafe_allow_html=True)
    st.markdown("""
        <h2 style='text-align: center; font-size: 1.5rem; font-weight: 400; margin-bottom: 0.875rem;'>
            Analyze and Enhance Your Resume with AI 
        </h2>
    """, unsafe_allow_html=True)
    
    # Centered flowchart toggle
    col_toggle, _, _ = st.columns([1, 1, 1])
    with col_toggle:
        show_flowchart = st.toggle('Show System Architecture', False)
    
    if show_flowchart:
        st.graphviz_chart(display_architecture_diagram())

    # Initialize variables
    resume_text = None
    job_description = None
    original_file = None

    # Side-by-side inputs with equal width
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("##### Upload Resume")
        uploaded_file = st.file_uploader(
            "Drop your resume file here",
            type="docx",
            label_visibility="collapsed"
        )
        
        if uploaded_file is not None:
            # Convert the docx file to text
            resume_text = docx2txt.process(uploaded_file)
            original_file = uploaded_file

    with col2:
        st.markdown("##### Job Description")
        job_description = st.text_area(
            "Paste job description",
            placeholder="Paste the job description here...",
            height=220,
            label_visibility="collapsed"
        )

    # Centered API key input first
    key = st.text_input(
            "GROQ API Key",
            type="password",
            placeholder="Enter your GROQ API key...",
            help="Your API key will not be stored"
        )

    if key:
        client = groq.Client(api_key=key)
        
        # Centered action buttons
        col_buttons, _, _ = st.columns([200,1,1])
        with col_buttons:
            col_b1, col_b2, col_b3 = st.columns(3)
            
            process_quick = col_b1.button("Quick Analysis")
            process_deep = col_b2.button("In-Depth Analysis")
            process_enhance = col_b3.button("Enhance Resume")

        #to check if any of the buttons are clicked
        if any([process_quick, process_deep, process_enhance]):
            if not resume_text:
                st.error("Please upload your resume.")
            elif not job_description:
                st.error("Please provide the job description.")
            else:
                try:
                    with st.spinner("Processing your resume..."):
                        if process_quick:
                            analysis = quick_analyze_resume(client, resume_text, job_description)
                            st.markdown("### Quick Analysis Results")
                            st.markdown(analysis)
                        
                        elif process_deep:
                            analysis = in_depth_analyze_resume(client, resume_text, job_description)
                            st.markdown("### Detailed Analysis")
                            st.markdown(analysis)
                        
                        else:  # Enhance Resume
                            analysis = in_depth_analyze_resume(client, resume_text, job_description)
                            enhanced_resume = enhance_resume(client, analysis, resume_text)
                            formatted_html = generate_formatted_resume_html(client, enhanced_resume)

                            # Results tabs
                            tab1, tab2, tab3 = st.tabs([
                                "Enhanced Resume",
                                "Compare Changes",
                                "Download Options"
                            ])
                            
                            with tab1:
                                # Single HTML component for the enhanced resume
                                st.components.v1.html(
                                    f"""
                                    <div class="resume-html">
                                        {formatted_html}
                                    </div>
                                    """,
                                    height=800,
                                    scrolling=True
                                )
                            
                            with tab2:
                                st.components.v1.html(
                                    show_diff(resume_text, enhanced_resume),
                                    height=600,
                                    scrolling=True
                                )
                            
                            with tab3:
                                col1, col2, col3 = st.columns(3)
                                with col1:
                                    if original_file:
                                        updated_docx = update_docx(original_file, enhanced_resume)
                                    else:
                                        updated_docx = BytesIO()
                                        doc = docx.Document()
                                        doc.add_paragraph(enhanced_resume)
                                        doc.save(updated_docx)
                                        updated_docx.seek(0)
                                        
                                    st.download_button(
                                        "📄 Download DOCX",
                                        data=updated_docx,
                                        file_name="enhanced_resume.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                    )
                                with col2:
                                    st.download_button(
                                        "🌐 Download HTML",
                                        data=formatted_html,
                                        file_name="enhanced_resume.html",
                                        mime="text/html"
                                    )
                                with col3:
                                    st.download_button(
                                        "📝 Download TXT",
                                        data=enhanced_resume,
                                        file_name="enhanced_resume.txt",
                                        mime="text/plain"
                                    )
                except groq.RateLimitError as e:
                    st.error("API rate limit exceeded. Please try again later or use a different API key.")
                except Exception as e:
                    st.error(f"An error occurred: {e}")
    else:
        st.info("👆 Please enter your GROQ API key to get started.")

if __name__ == "__main__":
    main()